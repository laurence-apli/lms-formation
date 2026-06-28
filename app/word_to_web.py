"""
Moteur d'import Word -> HTML, branché réellement sur le serveur (contrairement
à la maquette, qui ne pouvait que SIMULER cette zone -- un navigateur ne peut
pas exécuter de Python).

Même logique exacte que la version testée et validée à l'étape 1 du projet,
refactorisée en classe pour être réutilisable sans variables globales
(important : plusieurs imports peuvent arriver en même temps sur le serveur).

Règles appliquées, fidèles à ce qui a été validé avec Fabien :
 1. Tableau 1 colonne, texte gras+italique  -> encadré pull-quote doré
 2. Tableau 2 colonnes                       -> comparatif côte-à-côte
 3. Paragraphe de style Heading 1            -> titre de chapitre
 4. Paragraphe tout en MAJUSCULES + gras      -> sous-titre de section
 5. Paragraphe entièrement en italique        -> phrase d'accroche
 6. Liste à puces Word                        -> liste à puces HTML standard

Principe directeur, validé explicitement : "le Word pilote tout" -- si une
convention n'est pas respectée dans le fichier, le rendu reste simple, sans
qu'aucune mise en forme ne soit inventée par le système.
"""
import io
import base64
import html as htmlmod
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.table import Table

EXT_TO_MIME = {
    "png": "image/png", "jpeg": "image/jpeg", "jpg": "image/jpeg",
    "gif": "image/gif", "bmp": "image/bmp", "emf": "image/x-emf", "wmf": "image/x-wmf",
}

# Une vraie photo de contenu pèse largement plus que ce seuil ; en dessous,
# il s'agit presque toujours d'un gabarit décoratif (cadre, ombre, filigrane)
# que Word superpose à la vraie image -- ignoré volontairement (décision validée).
MIN_IMAGE_BYTES = 15000

MC_NS = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"

ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
}


def esc(s):
    return htmlmod.escape(s, quote=False)


class ImporteurWord:
    """Encapsule l'état nécessaire à l'import d'UN document, pour éviter tout
    risque de variables partagées entre deux imports simultanés sur le serveur."""

    def __init__(self, doc: Document):
        self.doc = doc

    # ---------- Images ----------

    def extract_image_as_data_uri(self, rel_id):
        try:
            part = self.doc.part.related_parts[rel_id]
            if len(part.blob) < MIN_IMAGE_BYTES:
                return None
            ext = (part.partname.ext or "png").lower()
            mime = EXT_TO_MIME.get(ext, "image/png")
            b64 = base64.b64encode(part.blob).decode("ascii")
            return f"data:{mime};base64,{b64}"
        except Exception:
            return None

    def find_images_in_paragraph(self, paragraph):
        found = []
        for ac in paragraph._p.iter(MC_NS + "AlternateContent"):
            choice = ac.find(MC_NS + "Choice")
            search_root = choice if choice is not None else ac
            for blip in search_root.iter(qn("a:blip")):
                rid = blip.get(qn("r:embed"))
                if rid:
                    data_uri = self.extract_image_as_data_uri(rid)
                    if data_uri:
                        found.append(data_uri)
        for drawing in paragraph._p.iter(qn("w:drawing")):
            parent = drawing.getparent()
            skip = False
            while parent is not None:
                if parent.tag == MC_NS + "AlternateContent":
                    skip = True
                    break
                parent = parent.getparent()
            if skip:
                continue
            for blip in drawing.iter(qn("a:blip")):
                rid = blip.get(qn("r:embed"))
                if rid:
                    data_uri = self.extract_image_as_data_uri(rid)
                    if data_uri:
                        found.append(data_uri)
        return found

    # ---------- Texte, gras/italique, taille, alignement ----------

    @staticmethod
    def runs_text_and_props(paragraph):
        text = paragraph.text.strip()
        if not text:
            return "", False, False
        all_bold = all(r.bold for r in paragraph.runs if r.text.strip()) if paragraph.runs else False
        all_italic = all(r.italic for r in paragraph.runs if r.text.strip()) if paragraph.runs else False
        return text, all_bold, all_italic

    @staticmethod
    def run_font_px(run):
        if run.font.size is not None:
            pt = run.font.size.pt
            return round(pt * 1.333, 1)
        return None

    def paragraph_dominant_font_px(self, paragraph):
        for r in paragraph.runs:
            if r.text.strip():
                px = self.run_font_px(r)
                if px:
                    return px
        return None

    @staticmethod
    def paragraph_align_css(paragraph):
        align = paragraph.alignment
        if align is None:
            return None
        return ALIGN_MAP.get(align)

    @staticmethod
    def is_uppercase_heading(text, bold):
        letters = [c for c in text if c.isalpha()]
        if not letters:
            return False
        return bold and text == text.upper() and len(text) > 3

    # ---------- Rendu d'un paragraphe ----------

    def render_paragraph_html(self, paragraph):
        images = self.find_images_in_paragraph(paragraph)
        img_html = "".join(f'<img class="chapter-image" src="{p}" alt="">' for p in images)

        text, bold, italic = self.runs_text_and_props(paragraph)
        if not text:
            return img_html

        style_name = paragraph.style.name if paragraph.style else ""
        font_px = self.paragraph_dominant_font_px(paragraph)
        align_css = self.paragraph_align_css(paragraph)

        style_parts = []
        if font_px:
            style_parts.append(f"font-size:{font_px}px")
        if align_css:
            style_parts.append(f"text-align:{align_css}")
        size_style = f' style="{";".join(style_parts)};"' if style_parts else ""

        if style_name.startswith("Heading") or style_name == "Title":
            return img_html + f'<h1 class="chapter-title"{size_style}>{esc(text)}</h1>'

        if self.is_uppercase_heading(text, bold):
            return img_html + f'<h2 class="section-title"{size_style}>{esc(text)}</h2>'

        if italic and not bold:
            return img_html + f'<p class="accent-line"{size_style}>{esc(text)}</p>'

        parts = []
        for r in paragraph.runs:
            t = esc(r.text)
            if not t:
                continue
            run_px = self.run_font_px(r)
            if r.bold:
                t = f"<strong>{t}</strong>"
            if r.italic:
                t = f"<em>{t}</em>"
            if run_px:
                t = f'<span style="font-size:{run_px}px;">{t}</span>'
            parts.append(t)
        inner = "".join(parts) if parts else esc(text)
        p_style = f' style="text-align:{align_css};"' if align_css else ""
        return img_html + f"<p{p_style}>{inner}</p>"

    # ---------- Listes ----------

    @staticmethod
    def render_bullet_list_block(paragraphs):
        items = "".join(f"<li>{esc(p.text.strip())}</li>" for p in paragraphs if p.text.strip())
        return f'<ul class="standard-list">{items}</ul>'

    @staticmethod
    def is_list_paragraph(paragraph):
        pPr = paragraph._p.find(qn("w:pPr"))
        if pPr is None:
            return False
        return pPr.find(qn("w:numPr")) is not None

    # ---------- Tableaux ----------

    @staticmethod
    def cell_text_props(cell):
        full_text = cell.text.strip()
        bolds, italics = [], []
        for p in cell.paragraphs:
            for r in p.runs:
                if r.text.strip():
                    bolds.append(bool(r.bold))
                    italics.append(bool(r.italic))
        all_bold = all(bolds) if bolds else False
        all_italic = all(italics) if italics else False
        return full_text, all_bold, all_italic

    def render_table_html(self, table):
        n_cols = len(table.columns)
        rows = table.rows

        if n_cols == 1:
            cell = rows[0].cells[0]
            text, bold, italic = self.cell_text_props(cell)
            if not text:
                return '<hr class="gold-rule-thin">'
            paras = [p for p in cell.paragraphs if p.text.strip()]
            label = None
            body_paras = []
            for p in paras:
                ptext, pbold, pital = self.runs_text_and_props(p)
                p_align = self.paragraph_align_css(p)
                if pbold and not pital and label is None and ptext == ptext.upper():
                    label = ptext
                else:
                    body_paras.append((ptext, p_align))
            label_html = f'<div class="pq-label">{esc(label)}</div>' if label else ""
            body_html = "".join(
                f'<p{f" style=\"text-align:{a};\"" if a else ""}>{esc(t)}</p>'
                for t, a in body_paras
            )
            return f'<div class="pull-quote">{label_html}<div class="pq-body">{body_html}</div></div>'

        if n_cols == 2:
            col_html = []
            css_classes = ["before", "after"]
            for idx, cell in enumerate(rows[0].cells):
                label = None
                items = []
                for p in cell.paragraphs:
                    ptext, pbold, pital = self.runs_text_and_props(p)
                    if not ptext:
                        continue
                    if label is None and pbold:
                        label = ptext
                    else:
                        items.append(ptext)
                li_html = "".join(f"<li>{esc(i)}</li>" for i in items)
                col_html.append(
                    f'<div class="compare-col {css_classes[idx % 2]}">'
                    f'<div class="ct-label">{esc(label or "")}</div>'
                    f"<ul>{li_html}</ul></div>"
                )
            return f'<div class="compare-grid">{"".join(col_html)}</div>'

        out = ['<table class="generic-table"><tbody>']
        for row in rows:
            out.append("<tr>" + "".join(f"<td>{esc(c.text)}</td>" for c in row.cells) + "</tr>")
        out.append("</tbody></table>")
        return "".join(out)

    # ---------- Parcours séquentiel complet du document ----------

    def importer(self) -> str:
        """Parcourt le document dans l'ordre réel (paragraphes ET tableaux),
        applique mécaniquement les 6 règles, renvoie le HTML du contenu --
        SANS la page <html> complète (ce contenu sera inséré dans le rendu
        de la plateforme, pas affiché comme une page indépendante)."""
        html_blocks = []
        pending_list = []
        body = self.doc.element.body

        def flush_list():
            if pending_list:
                html_blocks.append(self.render_bullet_list_block(pending_list))
                pending_list.clear()

        for child in body.iterchildren():
            tag = child.tag
            if tag == qn("w:p"):
                paragraph = Paragraph(child, self.doc)
                if self.is_list_paragraph(paragraph):
                    if paragraph.text.strip():
                        pending_list.append(paragraph)
                    continue
                flush_list()
                block = self.render_paragraph_html(paragraph)
                if block:
                    html_blocks.append(block)
            elif tag == qn("w:tbl"):
                flush_list()
                table = Table(child, self.doc)
                html_blocks.append(self.render_table_html(table))

        flush_list()
        return "\n".join(html_blocks)


def importer_word_depuis_bytes(contenu_fichier: bytes) -> str:
    """Point d'entrée principal pour le serveur : reçoit les octets bruts d'un
    fichier .docx uploadé, renvoie le HTML du contenu prêt à être stocké dans
    Chapitre.contenu_html (ou Module/Formation.presentation_html)."""
    doc = Document(io.BytesIO(contenu_fichier))
    importeur = ImporteurWord(doc)
    return importeur.importer()
